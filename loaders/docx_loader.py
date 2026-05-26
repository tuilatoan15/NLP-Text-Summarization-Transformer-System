"""DOCX loader using python-docx with Mammoth fallback for format preservation."""

from __future__ import annotations

import hashlib
import re
from pathlib import Path
from typing import Any

from docx import Document

from pipeline.schema import DocumentElement, DocumentMetadata, ExtractionConfig, ExtractedDocument
from preprocess.cleaner import normalize_unicode
from utils.logger import logger
from utils.metrics import extraction_quality_score


class DOCXLoader:
    """Extract paragraphs, headings, bullets, tables, and core metadata from DOCX."""

    def __init__(self, config: ExtractionConfig | None = None) -> None:
        self.config = config or ExtractionConfig()

    def load(self, path: str | Path) -> ExtractedDocument:
        file_path = Path(path)
        warnings: list[str] = []
        elements, metadata = self._extract_with_python_docx(file_path)
        text = self._join_elements(elements)
        quality = extraction_quality_score(text)

        if self.config.docx_use_mammoth_fallback and float(quality["score"]) < 0.35:
            fallback_elements = self._extract_with_mammoth(file_path)
            fallback_text = self._join_elements(fallback_elements)
            fallback_quality = extraction_quality_score(fallback_text)
            if float(fallback_quality["score"]) > float(quality["score"]):
                elements = fallback_elements
                text = fallback_text
                quality = fallback_quality
                metadata.extraction_engine = "mammoth"
                warnings.append("Used Mammoth fallback because python-docx extraction quality was low.")

        metadata.quality_score = float(quality["score"])
        metadata.language = str(quality["language"])
        structure = self._build_structure(elements)
        return ExtractedDocument(
            document_id=self._document_id(file_path, text),
            metadata=metadata,
            text=text,
            elements=elements,
            structure=structure,
            warnings=warnings,
        )

    def _extract_with_python_docx(self, path: Path) -> tuple[list[DocumentElement], DocumentMetadata]:
        document = Document(str(path))
        elements: list[DocumentElement] = []
        for paragraph in document.paragraphs:
            text = normalize_unicode(paragraph.text.strip())
            if not text:
                continue
            style = (paragraph.style.name or "").lower() if paragraph.style else ""
            element_type = self._classify_paragraph(text, style)
            elements.append(
                DocumentElement(
                    text=text,
                    element_type=element_type,
                    metadata={"style": style},
                    level=self._heading_level(style),
                )
            )

        if self.config.preserve_tables:
            for table_index, table in enumerate(document.tables):
                rows = []
                for row in table.rows:
                    cells = [normalize_unicode(cell.text.strip()) for cell in row.cells]
                    if any(cells):
                        rows.append(" | ".join(cells))
                if rows:
                    elements.append(
                        DocumentElement(
                            text="\n".join(rows),
                            element_type="table",
                            metadata={"table_index": table_index},
                        )
                    )

        props = document.core_properties
        metadata = DocumentMetadata(
            source_path=str(path),
            source_type="docx",
            title=props.title or path.stem,
            author=props.author or None,
            pages=None,
            created_at=props.created.isoformat() if props.created else None,
            modified_at=props.modified.isoformat() if props.modified else None,
            extraction_engine="python-docx",
            extra={
                "subject": props.subject,
                "keywords": props.keywords,
                "category": props.category,
                "comments": props.comments,
            },
        )
        return elements, metadata

    def _extract_with_mammoth(self, path: Path) -> list[DocumentElement]:
        try:
            import mammoth
            from bs4 import BeautifulSoup
        except Exception as exc:
            logger.debug("Mammoth unavailable: %s", exc)
            return []

        try:
            with path.open("rb") as handle:
                result = mammoth.convert_to_html(handle)
            soup = BeautifulSoup(result.value, "lxml")
            elements: list[DocumentElement] = []
            for node in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "table"]):
                text = normalize_unicode(node.get_text(separator=" ", strip=True))
                if not text:
                    continue
                if node.name and node.name.startswith("h"):
                    element_type = "heading"
                    level = int(node.name[1])
                elif node.name == "li":
                    element_type = "bullet"
                    level = None
                elif node.name == "table":
                    element_type = "table"
                    level = None
                else:
                    element_type = "paragraph"
                    level = None
                elements.append(DocumentElement(text=text, element_type=element_type, level=level))
            return elements
        except Exception as exc:
            logger.warning("Mammoth extraction failed for %s: %s", path, exc)
            return []

    @staticmethod
    def _classify_paragraph(text: str, style: str) -> str:
        if "heading" in style or "title" in style:
            return "heading"
        if "list" in style or re.match(r"^\s*(?:[-*+•]|\d+[.)])\s+", text):
            return "bullet"
        return "paragraph"

    @staticmethod
    def _heading_level(style: str) -> int | None:
        match = re.search(r"heading\s+(\d+)", style)
        return int(match.group(1)) if match else None

    @staticmethod
    def _build_structure(elements: list[DocumentElement]) -> dict[str, Any]:
        sections: list[dict[str, Any]] = []
        current_path: list[str] = []
        for idx, element in enumerate(elements):
            if element.element_type == "heading":
                level = element.level or 1
                current_path = current_path[: max(0, level - 1)] + [element.text]
                element.section_path = current_path.copy()
                sections.append(
                    {
                        "title": element.text,
                        "level": level,
                        "page": element.page_number,
                        "element_index": idx,
                        "path": current_path.copy(),
                    }
                )
            else:
                element.section_path = current_path.copy()
        return {"sections": sections, "tables": sum(1 for element in elements if element.element_type == "table")}

    @staticmethod
    def _join_elements(elements: list[DocumentElement]) -> str:
        return "\n\n".join(element.text.strip() for element in elements if element.text.strip())

    @staticmethod
    def _document_id(path: Path, text: str) -> str:
        digest = hashlib.sha1()
        digest.update(str(path.resolve()).encode("utf-8", errors="ignore"))
        digest.update(text[:8192].encode("utf-8", errors="ignore"))
        return digest.hexdigest()
